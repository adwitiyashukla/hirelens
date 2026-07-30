"""Readers that turn a file into a :class:`SourceDocument` with intact offsets.

The PDF reader is the interesting one. PyMuPDF's ``page.get_text()`` returns a
plain string and throws the geometry away, which would make citation highlights
impossible. So we walk the structured ``"dict"`` output instead, rebuild lines
ourselves, and record a bounding box for each. That costs about twenty extra
lines and buys the single most convincing feature in the demo.

Resumes are also frequently two-column. Reading a two-column page in raw
extraction order interleaves the columns and produces garbage like
"Education Python Stanford Docker". We detect the column split and read each
column top to bottom, which is what a human does.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from hirelens.ingest.document import (
    BoundingBox,
    IngestionError,
    SourceDocument,
    SourceFormat,
    TextAccumulator,
    detect_format,
)

logger = logging.getLogger(__name__)

# A line whose font is this much larger than the page median is treated as a
# section heading. Resumes are visually consistent enough for this to work.
_HEADING_FONT_RATIO = 1.15

# Two-column layouts are only worth handling when both columns carry real
# content. Below this share of lines we assume a single column with a stray
# right-aligned date.
_MIN_COLUMN_SHARE = 0.25


def read_document(path: str | Path) -> SourceDocument:
    """Read any supported file into a :class:`SourceDocument`."""
    path = Path(path)
    if not path.exists():
        raise IngestionError(f"File not found: {path}")
    if path.stat().st_size == 0:
        raise IngestionError(f"File is empty: {path}")

    fmt = detect_format(path)
    reader = {
        SourceFormat.PDF: read_pdf,
        SourceFormat.DOCX: read_docx,
        SourceFormat.TEXT: read_text,
    }[fmt]
    doc = reader(path)

    if doc.is_probably_scanned:
        logger.warning(
            "%s yielded only %d characters across %d page(s). This is almost "
            "certainly a scanned document with no text layer; OCR is needed.",
            path.name,
            doc.char_count,
            doc.page_count,
        )
    return doc


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def read_pdf(path: Path) -> SourceDocument:
    try:
        import pymupdf
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise IngestionError("PyMuPDF is required to read PDFs: pip install pymupdf") from exc

    payload = path.read_bytes()
    acc = TextAccumulator()

    with pymupdf.open(stream=payload, filetype="pdf") as pdf:
        page_count = pdf.page_count
        metadata = {k: str(v) for k, v in (pdf.metadata or {}).items() if v}

        for page_index in range(page_count):
            page = pdf[page_index]
            lines = _extract_page_lines(page)
            if not lines:
                continue

            lines = _order_lines(lines, page_width=float(page.rect.width))
            median_size = _median([ln["size"] for ln in lines if ln["size"]]) or 0.0

            for line in lines:
                acc.add_line(
                    line["text"],
                    page=page_index + 1,
                    bbox=line["bbox"],
                    is_heading=bool(
                        median_size and line["size"] >= median_size * _HEADING_FONT_RATIO
                    ),
                    font_size=line["size"],
                )
            acc.add_page_break()

    return acc.build(
        document_id=SourceDocument.make_id(payload),
        filename=path.name,
        source_format=SourceFormat.PDF,
        page_count=page_count,
        metadata=metadata,
    )


def _extract_page_lines(page: Any) -> list[dict[str, Any]]:
    """Pull one dict per visual line, carrying text, box and dominant font size."""
    raw = page.get_text("dict")
    lines: list[dict[str, Any]] = []

    for block in raw.get("blocks", []):
        # type 0 is text; type 1 is an image, which has no characters for us.
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans)
            if not text.strip():
                continue
            x0, y0, x1, y1 = line["bbox"]
            # Use the widest span's size: a line reading "PROJECTS  (2024)" should
            # be classified by its heading text, not the small parenthetical.
            dominant = max(spans, key=lambda s: len(s.get("text", "")), default={})
            lines.append(
                {
                    "text": text,
                    "bbox": BoundingBox(x0=x0, y0=y0, x1=x1, y1=y1),
                    "size": float(dominant.get("size", 0.0)),
                }
            )
    return lines


def _order_lines(lines: list[dict[str, Any]], *, page_width: float) -> list[dict[str, Any]]:
    """Sort lines into human reading order, handling two-column resumes.

    Single column: plain top-to-bottom. Two columns: everything in the left
    column top-to-bottom, then everything in the right column. Getting this wrong
    is the most common reason naive resume parsers produce nonsense.
    """
    midpoint = page_width / 2
    left = [ln for ln in lines if ln["bbox"].x0 < midpoint]
    right = [ln for ln in lines if ln["bbox"].x0 >= midpoint]

    share = min(len(left), len(right)) / max(len(lines), 1)
    if share < _MIN_COLUMN_SHARE:
        return sorted(lines, key=lambda ln: (round(ln["bbox"].y0, 1), ln["bbox"].x0))

    # Genuine two-column layout: read each side independently.
    logger.debug("two-column layout detected (%d left / %d right)", len(left), len(right))
    key = lambda ln: (round(ln["bbox"].y0, 1), ln["bbox"].x0)  # noqa: E731
    return sorted(left, key=key) + sorted(right, key=key)


def _median(values: list[float]) -> float | None:
    clean = sorted(v for v in values if v > 0)
    if not clean:
        return None
    mid = len(clean) // 2
    if len(clean) % 2:
        return clean[mid]
    return (clean[mid - 1] + clean[mid]) / 2


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def read_docx(path: Path) -> SourceDocument:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise IngestionError(
            "python-docx is required to read .docx: pip install python-docx"
        ) from exc

    payload = path.read_bytes()
    acc = TextAccumulator()
    document = docx.Document(str(path))

    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        acc.add_line(paragraph.text, page=1, is_heading="heading" in style or style == "title")

    # Resumes frequently keep skills and dates in tables, and dropping them loses
    # real signal, so flatten each row into a readable line.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                acc.add_line(" | ".join(cells), page=1)

    return acc.build(
        document_id=SourceDocument.make_id(payload),
        filename=path.name,
        source_format=SourceFormat.DOCX,
        page_count=1,
    )


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def read_text(path: Path) -> SourceDocument:
    payload = path.read_bytes()
    content = payload.decode("utf-8", errors="replace")
    acc = TextAccumulator()

    for line in content.splitlines():
        # A short all-caps line in a plain-text resume is nearly always a header.
        stripped = line.strip()
        is_heading = bool(stripped) and stripped.isupper() and len(stripped) < 40
        acc.add_line(line, page=1, is_heading=is_heading)

    return acc.build(
        document_id=SourceDocument.make_id(payload),
        filename=path.name,
        source_format=SourceFormat.TEXT,
        page_count=1,
    )
